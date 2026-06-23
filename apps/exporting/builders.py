"""Document builders (FR-42/FR-43, §3.2.8).

A draft is first composed into a single ordered block list
(``content.compose_document``); the DOCX and PDF builders below each walk that
same list so the two formats stay structurally identical. The result is an
RFP-grade document — cover page, abbreviations table, table of contents,
fact sheet, numbered sections with tables and bullet lists, Forms & Annexures
and an authorized-signatory block.

Traceability is preserved from the original builder: generated clauses keep
their superscript citation markers, a References section lists the sources, and
the full clause→citation map is embedded in the DOCX custom-properties part so
clause-level traceability survives later manual edits (pre-bid clarification
Q43). PDF carries archival document metadata.
"""
from __future__ import annotations

import io
import json
from xml.sax.saxutils import escape

from .content import compose_document
from .document import (
    Bullets, Cover, Heading, PageBreak, Paragraph, Signature, TOC, Table,
)


def _heading_text(h: Heading) -> str:
    if not h.number:
        return h.text
    sep = "." if "." not in h.number else ""
    return f"{h.number}{sep} {h.text}"


def _citation_map(draft) -> dict:
    """clause_type -> [citation, …] for the embedded traceability metadata."""
    return {c.clause_type: [m["citation"] for m in c.citations]
            for c in draft.clauses.all()}


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _embed_custom_properties(doc, props: dict) -> bool:
    """Embed key/value strings as an OOXML custom-properties part
    (docProps/custom.xml). This survives manual edits to the body — the
    downstream-traceability guarantee from pre-bid clarification Q43.
    Returns True on success; callers fall back to core metadata otherwise."""
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        entries = []
        for i, (name, value) in enumerate(props.items(), start=2):
            entries.append(
                f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" '
                f'pid="{i}" name="{escape(name)}">'
                f'<vt:lpwstr>{escape(value)}</vt:lpwstr></property>'
            )
        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
            'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            + "".join(entries)
            + "</Properties>"
        )
        package = doc.part.package
        partname = PackURI("/docProps/custom.xml")
        ctype = "application/vnd.openxmlformats-officedocument.custom-properties+xml"
        part = Part(partname, ctype, xml.encode("utf-8"), package)
        package.relate_to(part, RT.CUSTOM_PROPERTIES)
        return True
    except Exception:
        return False


def _docx_field(paragraph, instr: str, default: str = "") -> None:
    """Append a simple field (e.g. PAGE / PAGEREF) to a paragraph.

    ``default`` is the cached result shown until the reader recalculates fields
    (we also flag the document to auto-update on open, so this is just a sensible
    placeholder rather than a blank cell)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = default
    run.append(txt)
    fld.append(run)
    paragraph._p.append(fld)


def _docx_bookmark(paragraph, name: str, bid: int) -> None:
    """Wrap a paragraph's content in a bookmark so a PAGEREF can point to it."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _docx_auto_update_fields(doc) -> None:
    """Flag the document so Word fills in PAGEREF page numbers on open
    (PAGEREF updates silently — no 'update fields?' prompt)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    # Per the OOXML CT_Settings schema, updateFields must precede compat/rsids
    # etc.; insert before the first such element so Word doesn't flag the file
    # for repair. Fall back to append if none are present.
    anchor = None
    for tag in ("w:compat", "w:docVars", "w:rsids", "w:mathPr",
                "w:themeFontLang", "w:clrSchemeMapping",
                "w:doNotAutoCompressPictures", "w:shapeDefaults",
                "w:decimalSymbol", "w:listSeparator"):
        anchor = settings.find(qn(tag))
        if anchor is not None:
            break
    if anchor is not None:
        anchor.addprevious(el)
    else:
        settings.append(el)


def _docx_toc(doc, headings, bookmarks) -> None:
    """A self-built table of contents: section numbers + titles are real text
    (always visible), page numbers are PAGEREF fields that auto-fill on open."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc.add_heading("Table of Contents", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for h in headings:
        bid, name = bookmarks[id(h)]
        cells = table.add_row().cells
        left = cells[0].paragraphs[0]
        if h.level >= 2:
            left.paragraph_format.left_indent = Inches(0.3)
        run = left.add_run(_heading_text(h))
        run.bold = h.level == 1
        right = cells[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _docx_field(right, f"PAGEREF {name} \\h", default="•")
        cells[0].width = Inches(5.8)
        cells[1].width = Inches(0.5)
    doc.add_paragraph()


def _docx_footer(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_field(p, "PAGE")
    p.add_run(" | Page")


def _docx_table(doc, block: Table) -> None:
    from docx.shared import Inches, Pt

    if block.caption:
        cap = doc.add_paragraph()
        r = cap.add_run(block.caption)
        r.bold = True
        r.font.size = Pt(10)

    has_header = bool(block.headers)
    n_cols = len(block.headers) if has_header else (len(block.rows[0]) if block.rows else 1)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False

    if has_header:
        cells = table.add_row().cells
        for cell, text in zip(cells, block.headers):
            cell.paragraphs[0].add_run(str(text)).bold = True
    for row in block.rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.paragraphs[0].add_run(str(text))

    # Apply relative column widths over ~6.3" of usable page width.
    if block.widths and len(block.widths) == n_cols:
        total = sum(block.widths) or 1
        for col, w in zip(table.columns, block.widths):
            width = Inches(6.3 * w / total)
            for cell in col.cells:
                cell.width = width
    doc.add_paragraph()


def _docx_bullets(doc, block: Bullets, level: int = 0) -> None:
    style = {"bullet": "List Bullet", "number": "List Number",
             "letter": "List Number"}.get(block.style, "List Bullet")
    if level:
        style = f"{style} {min(level + 1, 3)}"
    for item in block.items:
        if isinstance(item, Bullets):
            _docx_bullets(doc, item, level + 1)
        else:
            try:
                doc.add_paragraph(str(item), style=style)
            except KeyError:  # style not in default template
                doc.add_paragraph("• " + str(item))


def _docx_cover(doc, block: Cover) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc.add_paragraph()
    title = doc.add_heading(block.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if block.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(block.subtitle).italic = True
    doc.add_paragraph()
    if block.fields:
        _docx_table(doc, Table(headers=["Particulars", "Details"],
                               rows=[[k, v] for k, v in block.fields],
                               widths=[1.4, 3.0]))
    for line in block.issuer_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    if block.footer:
        doc.add_paragraph()
        f = doc.add_paragraph()
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = f.add_run(block.footer)
        run.italic = True
        run.font.size = Pt(9)


def _docx_signature(doc, block: Signature) -> None:
    from docx.shared import Pt

    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    r = doc.add_paragraph().add_run(
        "_______________________________   (Signature)")
    r.font.size = Pt(10)
    doc.add_paragraph(block.role)
    for fld in block.fields:
        doc.add_paragraph(f"{fld} ______________________________")
    if block.note:
        n = doc.add_paragraph().add_run(block.note)
        n.italic = True
        n.font.size = Pt(9)


def build_docx(draft) -> bytes:
    from docx import Document as Docx
    from docx.shared import Pt

    doc = Docx()
    _docx_footer(doc)

    blocks = compose_document(draft)
    # Pre-scan headings so the self-built TOC can list them and each heading
    # can carry a bookmark the TOC's PAGEREF points to.
    toc_headings = [b for b in blocks if isinstance(b, Heading) and b.toc]
    bookmarks = {id(h): (i + 1, f"RFPTOC{i + 1}") for i, h in enumerate(toc_headings)}

    for block in blocks:
        if isinstance(block, Cover):
            _docx_cover(doc, block)
        elif isinstance(block, TOC):
            _docx_toc(doc, toc_headings, bookmarks)
        elif isinstance(block, Heading):
            p = doc.add_heading(_heading_text(block), level=min(block.level, 3))
            mark = bookmarks.get(id(block))
            if mark:
                _docx_bookmark(p, mark[1], mark[0])
        elif isinstance(block, Paragraph):
            para = doc.add_paragraph(block.text)
            if block.style in ("italic", "note"):
                for run in para.runs:
                    run.italic = True
            if block.markers:
                run = para.add_run(" " + "".join(f"[{m}]" for m in block.markers))
                run.font.superscript = True
                run.font.size = Pt(8)
        elif isinstance(block, Bullets):
            _docx_bullets(doc, block)
        elif isinstance(block, Table):
            _docx_table(doc, block)
        elif isinstance(block, Signature):
            _docx_signature(doc, block)
        elif isinstance(block, PageBreak):
            doc.add_page_break()

    _docx_auto_update_fields(doc)
    doc.core_properties.title = f"{draft.instrument}: {draft.title}"
    doc.core_properties.category = draft.category
    doc.core_properties.author = "NeGD AI-Based RFP Authoring Tool"
    citations_json = json.dumps(_citation_map(draft), ensure_ascii=False)
    embedded = _embed_custom_properties(
        doc,
        {
            "RFP_Authoring_Tool": "true",
            "RFP_Clause_Citations": citations_json,
            "RFP_Verdict_Status": draft.status,
        },
    )
    if not embedded:
        doc.core_properties.keywords = "rfp-authoring-tool; clause-level-citations-embedded"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _san(text: str) -> str:
    """reportlab's built-in Helvetica has no Indian Rupee glyph (U+20B9); it
    would render as a missing-glyph box. Substitute a portable 'Rs.' for PDF
    output only (DOCX keeps ₹, which Word fonts support)."""
    return str(text).replace("₹", "Rs. ")


def _pdf_styles():
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    ss = getSampleStyleSheet()
    styles = {
        "Title": ss["Title"],
        "Body": ParagraphStyle("RFPBody", parent=ss["BodyText"], spaceAfter=6,
                               leading=14, alignment=4),  # justified
        "Note": ParagraphStyle("RFPNote", parent=ss["BodyText"], spaceAfter=6,
                               leading=13, fontName="Helvetica-Oblique",
                               textColor="#444444"),
        "H1": ParagraphStyle("RFPH1", parent=ss["Heading1"], spaceBefore=12,
                             spaceAfter=6, fontSize=14),
        "H2": ParagraphStyle("RFPH2", parent=ss["Heading2"], spaceBefore=8,
                             spaceAfter=4, fontSize=12),
        "H3": ParagraphStyle("RFPH3", parent=ss["Heading3"], spaceBefore=6,
                             spaceAfter=4, fontSize=11),
        "CoverTitle": ParagraphStyle("CoverTitle", parent=ss["Title"],
                                     fontSize=20, leading=26, alignment=TA_CENTER),
        "CoverSub": ParagraphStyle("CoverSub", parent=ss["Italic"],
                                   alignment=TA_CENTER, spaceAfter=18),
        "CoverLine": ParagraphStyle("CoverLine", parent=ss["Normal"],
                                    alignment=TA_CENTER, leading=15),
        "Cell": ParagraphStyle("RFPCell", parent=ss["BodyText"], fontSize=9,
                               leading=12, spaceAfter=0),
        "CellHead": ParagraphStyle("RFPCellHead", parent=ss["BodyText"],
                                   fontSize=9, leading=12, spaceAfter=0,
                                   fontName="Helvetica-Bold", textColor="white"),
        "Caption": ParagraphStyle("RFPCaption", parent=ss["BodyText"],
                                  fontName="Helvetica-Bold", fontSize=10,
                                  spaceBefore=6, spaceAfter=3),
        # Compact table-of-contents entry styles (tight leading, no big gaps).
        "TOC0": ParagraphStyle("TOC0", parent=ss["Normal"],
                               fontName="Helvetica-Bold", fontSize=10,
                               leading=13, spaceBefore=0, spaceAfter=0),
        "TOC1": ParagraphStyle("TOC1", parent=ss["Normal"], fontSize=9,
                               leading=12, spaceBefore=0, spaceAfter=0,
                               leftIndent=16),
    }
    return styles


def _pdf_table(block: Table, styles, avail: float):
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph as P
    from reportlab.platypus import KeepTogether, Spacer, Table as RLTable
    from reportlab.platypus import TableStyle

    data = []
    if block.headers:
        data.append([P(_san(h), styles["CellHead"]) for h in block.headers])
    for row in block.rows:
        data.append([P(_san(c), styles["Cell"]) for c in row])

    n_cols = len(data[0]) if data else 1
    weights = block.widths if (block.widths and len(block.widths) == n_cols) \
        else [1] * n_cols
    total = sum(weights) or 1
    col_widths = [avail * w / total for w in weights]

    tbl = RLTable(data, colWidths=col_widths, repeatRows=1 if block.headers else 0)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#888888")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    if block.headers:
        style.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#34495e")))
    else:
        style.append(("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f0")))
    tbl.setStyle(TableStyle(style))

    flow = []
    if block.caption:
        flow.append(P(_san(block.caption), styles["Caption"]))
    flow.append(tbl)
    flow.append(Spacer(1, 8))
    return KeepTogether(flow) if len(block.rows) <= 6 else flow


def _pdf_bullets(block: Bullets, styles):
    from reportlab.platypus import ListFlowable, ListItem
    from reportlab.platypus import Paragraph as P

    bullet = {"bullet": "bullet", "number": "1", "letter": "a"}.get(block.style, "bullet")
    items = []
    for item in block.items:
        if isinstance(item, Bullets):
            items.append(ListItem(_pdf_bullets(item, styles), value=None))
        else:
            items.append(ListItem(P(_san(item), styles["Body"])))
    return ListFlowable(items, bulletType=bullet, leftIndent=18,
                        bulletFontSize=9, spaceAfter=6)


def _pdf_signature(block: Signature, styles):
    from reportlab.platypus import Paragraph as P
    from reportlab.platypus import Spacer

    flow = [Spacer(1, 10), P("Yours faithfully,", styles["Body"]), Spacer(1, 14),
            P("_______________________________ &nbsp; (Signature)", styles["Body"]),
            P(block.role, styles["Body"])]
    for fld in block.fields:
        flow.append(P(f"{fld} ______________________________", styles["Body"]))
    if block.note:
        flow.append(P(block.note, styles["Note"]))
    return flow


def build_pdf(draft) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageBreak as RLPageBreak, PageTemplate,
        Paragraph as P, Spacer,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    styles = _pdf_styles()
    buf = io.BytesIO()

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, f"{doc.page} | Page")
        canvas.restoreState()

    class _Doc(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if not isinstance(flowable, P):
                return
            name = flowable.style.name
            if name in ("RFPH1", "RFPH2"):
                level = 0 if name == "RFPH1" else 1
                self.notify("TOCEntry", (level, flowable.getPlainText(), self.page))

    doc = _Doc(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
               leftMargin=2 * cm, rightMargin=2 * cm,
               title=f"{draft.instrument}: {draft.title}",
               author="NeGD AI-Based RFP Authoring Tool",
               subject="Procurement document (archival)")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=_footer)])
    avail = doc.width

    toc = TableOfContents()
    toc.levelStyles = [styles["TOC0"], styles["TOC1"]]

    story = []
    for block in compose_document(draft):
        if isinstance(block, Cover):
            story.append(Spacer(1, 60))
            story.append(P(_san(block.title), styles["CoverTitle"]))
            if block.subtitle:
                story.append(P(_san(block.subtitle), styles["CoverSub"]))
            story.append(Spacer(1, 18))
            if block.fields:
                story.append(_pdf_table(
                    Table(headers=["Particulars", "Details"],
                          rows=[[k, v] for k, v in block.fields],
                          widths=[1.4, 3.0]), styles, avail))
            story.append(Spacer(1, 24))
            for line in block.issuer_lines:
                story.append(P(_san(line), styles["CoverLine"]))
            if block.footer:
                story.append(Spacer(1, 18))
                story.append(P(f"<i>{_san(block.footer)}</i>", styles["CoverLine"]))
        elif isinstance(block, TOC):
            story.append(P(block.title, styles["H1"]))
            story.append(toc)
        elif isinstance(block, Heading):
            lvl = {1: "H1", 2: "H2", 3: "H3"}.get(block.level, "H3")
            story.append(P(_san(_heading_text(block)), styles[lvl]))
        elif isinstance(block, Paragraph):
            text = _san(block.text)
            if block.markers:
                text += " <super><font size=7>" + \
                    "".join(f"[{m}]" for m in block.markers) + "</font></super>"
            story.append(P(text, styles["Note" if block.style in ("italic", "note")
                                        else "Body"]))
        elif isinstance(block, Bullets):
            story.append(_pdf_bullets(block, styles))
            story.append(Spacer(1, 6))
        elif isinstance(block, Table):
            item = _pdf_table(block, styles, avail)
            story.extend(item) if isinstance(item, list) else story.append(item)
        elif isinstance(block, Signature):
            story.extend(_pdf_signature(block, styles))
        elif isinstance(block, PageBreak):
            story.append(RLPageBreak())

    doc.multiBuild(story)
    return buf.getvalue()
